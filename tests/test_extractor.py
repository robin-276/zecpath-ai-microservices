import os
import pytest
import docx
from reportlab.pdfgen import canvas
from parsers.extractor import ResumeExtractor

# --- Fixtures to generate dummy files for testing ---

@pytest.fixture
def dummy_pdf(tmp_path):
    file_path = tmp_path / "test_resume.pdf"
    c = canvas.Canvas(str(file_path))
    c.drawString(100, 750, "Experience")
    c.drawString(100, 730, "• Built Zecpath AI")
    c.drawString(100, 710, "• Python Developer")
    c.save()
    return str(file_path)

@pytest.fixture
def dummy_docx(tmp_path):
    file_path = tmp_path / "test_resume.docx"
    doc = docx.Document()
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, FastAPI, Machine Learning")
    doc.save(file_path)
    return str(file_path)

# --- Test Cases ---

def test_pdf_extraction(dummy_pdf):
    extractor = ResumeExtractor(dummy_pdf)
    raw_text = extractor.extract_raw_text()
    assert "Experience" in raw_text
    assert "Zecpath" in raw_text

def test_docx_extraction(dummy_docx):
    extractor = ResumeExtractor(dummy_docx)
    raw_text = extractor.extract_raw_text()
    assert "FastAPI" in raw_text

def test_cleaning_and_normalization():
    extractor = ResumeExtractor("dummy.pdf") # File doesn't matter here
    messy_text = "Experience\n\n\n\n\u2022 Python Developer    \n\u2022 AI Engineer  @@!!"
    clean_text = extractor.clean_and_normalize(messy_text)
    
    # Check bullet normalization
    assert "-" in clean_text
    assert "\u2022" not in clean_text
    
    # Check whitespace and noise normalization
    assert "  " not in clean_text
    assert "@@!!" not in clean_text
    
    # Check heading normalization
    assert "EXPERIENCE:" in clean_text

def test_process_and_store(dummy_docx, tmp_path):
    output_dir = tmp_path / "data"
    extractor = ResumeExtractor(dummy_docx)
    
    result = extractor.process_and_store(output_dir=str(output_dir))
    
    assert "cleaned_text" in result
    assert os.path.exists(os.path.join(str(output_dir), "test_resume_extracted.json"))